from __future__ import annotations

import hashlib
import math
import re
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document
import fitz

from backend.app.models import ChunkStrategy


SECTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("Abstract", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:abstract|摘要)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Introduction", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:introduction|引言|绪论)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Methods", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:methods?|methodology|方法|实验方法|研究方法)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Results", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:results?|结果|实验结果|研究结果)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Discussion", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:discussion|讨论|分析与讨论)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Limitations", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:limitations?|研究局限|局限性|局限与不足|研究不足|不足与展望|不足)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("FutureWork", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:future\s+work|future\s+research|未来研究|后续研究|研究展望|展望)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("Conclusion", re.compile(r"^\s*(?:(?:第?[一二三四五六七八九十]+[章节、.．]|\d+(?:\.\d+)*[.、．]?)\s*)?(?:conclusions?|结论与展望|结论|总结)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
    ("References", re.compile(r"^\s*(?:references|参考文献)\s*(?:[:：]|\s|$)", re.IGNORECASE)),
]

TOKEN_UNIT_PATTERN = re.compile(
    r"\s+|[\u4e00-\u9fff]|[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*|.",
    re.DOTALL,
)


@dataclass
class PageText:
    page: int
    text: str


@dataclass
class TextBlock:
    page: int
    text: str
    char_start: int
    char_end: int
    section: str
    block_type: str = "text"


@dataclass
class ChunkRecord:
    chunk_id: str
    document_id: str
    paper_name: str
    page: int
    section: str
    source: str
    file_hash: str
    text: str
    char_start: int
    char_end: int
    quote: str
    page_start: int
    page_end: int
    token_count: int
    chunk_type: str = "text"
    parent_id: str = ""
    parent_text: str = ""
    parent_page_start: int = 0
    parent_page_end: int = 0
    parent_char_start: int = 0
    parent_char_end: int = 0
    parent_token_count: int = 0
    image_id: str = ""
    image_path: str = ""
    bbox_json: str = ""


def compute_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def safe_upload_name(file_hash: str, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower() or ".pdf"
    return f"{file_hash[:16]}{suffix}"


def extract_pdf_text(pdf_path: Path) -> list[PageText]:
    pages: list[PageText] = []
    with fitz.open(pdf_path) as doc:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            pages.append(PageText(page=index, text=text))
    return pages


def extract_docx_text(docx_path: Path) -> list[PageText]:
    document = Document(docx_path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            blocks.append("\n".join(table_rows))

    pages: list[PageText] = []
    if not blocks:
        return pages

    buffer: list[str] = []
    page_number = 1
    token_count = 0
    for block in blocks:
        block_tokens = count_tokens(block)
        if buffer and token_count + block_tokens > 1200:
            pages.append(PageText(page=page_number, text="\n\n".join(buffer)))
            page_number += 1
            buffer = []
            token_count = 0
        buffer.append(block)
        token_count += block_tokens

    if buffer:
        pages.append(PageText(page=page_number, text="\n\n".join(buffer)))

    return pages


def extract_document_text(path: Path) -> list[PageText]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    raise ValueError("暂不支持这个文件类型。请上传 PDF 或 DOCX。")


def analyze_pages(pages: list[PageText]) -> dict[str, int | str | float]:
    text = "\n".join(page.text for page in pages if page.text)
    char_count = len(text)
    token_count = count_tokens(text)
    paragraph_count = len([p for p in re.split(r"\n\s*\n|(?<=。)\s+|(?<=\.)\s+", text) if p.strip()])
    cjk_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    latin_chars = len(re.findall(r"[A-Za-z]", text))
    formula_hits = len(re.findall(r"[\u03b1-\u03c9\u0391-\u03a9=+\-*/<>≤≥∑∫√]|p\s*[<=>]", text))
    table_hits = len(re.findall(r"\b(table|表)\s*\d+|\|", text, flags=re.IGNORECASE))
    page_count = len(pages)

    if cjk_chars > latin_chars * 1.5:
        language = "zh"
    elif latin_chars > cjk_chars * 1.5:
        language = "en"
    else:
        language = "mixed"

    return {
        "page_count": page_count,
        "char_count": char_count,
        "token_count": token_count,
        "paragraph_count": paragraph_count,
        "language": language,
        "formula_table_ratio": (formula_hits + table_hits * 2) / max(char_count, 1),
        "avg_chars_per_page": char_count / max(page_count, 1),
        "avg_tokens_per_page": token_count / max(page_count, 1),
    }


def choose_chunk_strategy(pages: list[PageText]) -> ChunkStrategy:
    stats = analyze_pages(pages)
    page_count = int(stats["page_count"])
    char_count = int(stats["char_count"])
    token_count = int(stats["token_count"])
    paragraph_count = int(stats["paragraph_count"])
    language = str(stats["language"])
    formula_table_ratio = float(stats["formula_table_ratio"])
    avg_tokens_per_page = float(stats["avg_tokens_per_page"])

    chunk_size = 720
    overlap = 90
    parent_chunk_size = 1350
    parent_overlap = 160
    splitter = "token-paragraph-cross-page-hybrid"
    reasons: list[str] = ["切分长度已按 token 控制，而不是按字符数控制。"]

    if page_count >= 35 or token_count >= 70_000:
        chunk_size = 950
        overlap = 140
        parent_chunk_size = 1800
        parent_overlap = 220
        reasons.append("文档较长，使用更大的 token chunk 减少跨页主题被切散。")
    elif page_count <= 8 and token_count <= 18_000:
        chunk_size = 520
        overlap = 70
        parent_chunk_size = 1050
        parent_overlap = 120
        reasons.append("文档较短，使用较小 token chunk 提高定位精度。")

    density = paragraph_count / max(page_count, 1)
    if density >= 18:
        chunk_size = max(460, chunk_size - 120)
        overlap = max(60, overlap - 20)
        reasons.append("段落密度高，降低 token chunk 大小以避免一个片段混入过多论点。")
    elif density <= 5 and avg_tokens_per_page > 900:
        chunk_size += 120
        overlap += 35
        parent_chunk_size += 180
        reasons.append("段落偏长，增加 overlap 保留上下文连续性。")

    if language == "zh":
        chunk_size = int(chunk_size * 0.9)
        overlap = int(overlap * 0.9)
        parent_chunk_size = int(parent_chunk_size * 0.9)
        reasons.append("中文信息密度较高，适当缩小 token chunk。")
    elif language == "mixed":
        reasons.append("检测到中英文混合，采用页面、段落和 token 混合切分。")

    if formula_table_ratio > 0.012:
        splitter = "layout-aware-token-page-hybrid"
        overlap = max(overlap, 130)
        parent_overlap = max(parent_overlap, 220)
        reasons.append("公式或表格比例较高，单独识别表格块并增加 overlap。")

    return ChunkStrategy(
        chunk_size=chunk_size,
        overlap=overlap,
        splitter=splitter,
        language=language,
        page_count=page_count,
        paragraph_count=paragraph_count,
        char_count=char_count,
        token_count=token_count,
        size_unit="tokens",
        parent_chunk_size=parent_chunk_size,
        parent_overlap=parent_overlap,
        reasons=reasons,
    )


def split_pages_into_chunks(
    *,
    pages: list[PageText],
    document_id: str,
    paper_name: str,
    file_hash: str,
    source: str,
    strategy: ChunkStrategy,
) -> list[ChunkRecord]:
    blocks = build_text_blocks(pages)
    chunks: list[ChunkRecord] = []
    index = 0
    buffer: list[TextBlock] = []

    def flush_buffer() -> None:
        nonlocal index, buffer
        if not buffer:
            return
        text = "\n\n".join(block.text for block in buffer)
        index = _append_chunk(
            chunks=chunks,
            index=index,
            document_id=document_id,
            paper_name=paper_name,
            page_start=buffer[0].page,
            page_end=buffer[-1].page,
            section=buffer[0].section,
            source=source,
            file_hash=file_hash,
            text=text,
            char_start=buffer[0].char_start,
            char_end=buffer[-1].char_end,
            chunk_type="text",
        )
        buffer = []

    block_index = 0
    while block_index < len(blocks):
        block = blocks[block_index]
        if block.block_type == "table":
            flush_buffer()
            table_blocks = [block]
            table_end = block_index + 1
            while table_end < len(blocks) and blocks[table_end].block_type == "table":
                table_blocks.append(blocks[table_end])
                table_end += 1

            context_blocks: list[TextBlock] = []
            if block_index > 0 and _is_table_context_block(blocks[block_index - 1]):
                context_blocks.append(blocks[block_index - 1])
            context_blocks.extend(table_blocks)
            if table_end < len(blocks) and _is_table_context_block(blocks[table_end]):
                context_blocks.append(blocks[table_end])

            table_text = "\n\n".join(item.text for item in context_blocks)
            first = context_blocks[0]
            last = context_blocks[-1]
            if count_tokens(table_text) > strategy.chunk_size:
                for piece, start_offset in split_long_text(table_text, strategy.chunk_size, strategy.overlap):
                    index = _append_chunk(
                        chunks=chunks,
                        index=index,
                        document_id=document_id,
                        paper_name=paper_name,
                        page_start=first.page,
                        page_end=last.page,
                        section=block.section,
                        source=source,
                        file_hash=file_hash,
                        text=piece,
                        char_start=first.char_start + start_offset,
                        char_end=first.char_start + start_offset + len(piece),
                        chunk_type="table",
                    )
            else:
                index = _append_chunk(
                    chunks=chunks,
                    index=index,
                    document_id=document_id,
                    paper_name=paper_name,
                    page_start=first.page,
                    page_end=last.page,
                    section=block.section,
                    source=source,
                    file_hash=file_hash,
                    text=table_text,
                    char_start=first.char_start,
                    char_end=last.char_end,
                    chunk_type="table",
                )
            block_index = table_end
            continue

        if buffer and _should_flush_for_block(buffer[-1], block):
            flush_buffer()

        candidate_blocks = [*buffer, block]
        candidate = "\n\n".join(item.text for item in candidate_blocks)
        if count_tokens(candidate) <= strategy.chunk_size:
            buffer.append(block)
            block_index += 1
            continue

        flush_buffer()
        if count_tokens(block.text) > strategy.chunk_size:
            for piece, start_offset in split_long_text(block.text, strategy.chunk_size, strategy.overlap):
                index = _append_chunk(
                    chunks=chunks,
                    index=index,
                    document_id=document_id,
                    paper_name=paper_name,
                    page_start=block.page,
                    page_end=block.page,
                    section=block.section,
                    source=source,
                    file_hash=file_hash,
                    text=piece,
                    char_start=block.char_start + start_offset,
                    char_end=block.char_start + start_offset + len(piece),
                    chunk_type="text",
                )
        else:
            buffer.append(block)
        block_index += 1

    flush_buffer()
    _add_neighbor_overlap(chunks, strategy.overlap)
    _attach_parent_contexts(chunks, strategy.parent_chunk_size)
    return chunks


def build_text_blocks(pages: list[PageText]) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    current_section = "Unknown"
    for page in pages:
        page_text = page.text.strip()
        if not page_text:
            continue
        current_section = detect_section(page_text, current_section)
        paragraphs = split_paragraphs(page_text)
        cursor = 0
        for paragraph in paragraphs:
            paragraph_start = page_text.find(paragraph, cursor)
            if paragraph_start < 0:
                paragraph_start = cursor
            paragraph_end = paragraph_start + len(paragraph)
            cursor = paragraph_end
            paragraph_section = detect_section(paragraph, current_section)
            current_section = paragraph_section
            blocks.append(
                TextBlock(
                    page=page.page,
                    text=paragraph,
                    char_start=paragraph_start,
                    char_end=paragraph_end,
                    section=paragraph_section,
                    block_type="table" if is_table_like_block(paragraph) else "text",
                )
            )
    return blocks


def split_paragraphs(text: str) -> list[str]:
    rough = re.split(r"\n\s*\n", text)
    paragraphs: list[str] = []
    for item in rough:
        lines = [line.strip() for line in item.splitlines() if line.strip()]
        if not lines:
            continue
        if len(lines) >= 2 and any(is_table_like_block(line) for line in lines):
            paragraphs.append("\n".join(lines))
        else:
            paragraphs.append(" ".join(lines))
    if len(paragraphs) <= 1:
        paragraphs = [part.strip() for part in re.split(r"(?<=[。.!?])\s+", text) if part.strip()]
    return paragraphs or [text]


def detect_section(text: str, fallback: str) -> str:
    first_lines = [line.strip() for line in text.splitlines()[:10] if line.strip()]
    for line in first_lines:
        for name, pattern in SECTION_PATTERNS:
            if pattern.search(line):
                return name
    return fallback


def is_table_like_block(text: str) -> bool:
    normalized = " ".join(text.split())
    if not normalized:
        return False
    if re.search(r"^(?:table|tab\.|表)\s*\d+|^表[一二三四五六七八九十]+", normalized, re.IGNORECASE):
        return True
    pipe_count = normalized.count("|")
    if pipe_count >= 2:
        return True
    lines = [line for line in text.splitlines() if line.strip()]
    digit_ratio = len(re.findall(r"\d", normalized)) / max(len(normalized), 1)
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", normalized))
    repeated_spacing_rows = sum(1 for line in lines if len(re.split(r"\s{2,}", line.strip())) >= 3)
    return (
        len(normalized) > 80
        and (
            repeated_spacing_rows >= 2
            or (digit_ratio > 0.24 and cjk_count < 100 and len(lines) >= 2)
        )
    )


def split_long_text(text: str, chunk_size: int, overlap: int) -> list[tuple[str, int]]:
    sentences = split_sentences(text)
    if len(sentences) > 1:
        pieces: list[tuple[str, int]] = []
        buffer = ""
        buffer_start = 0
        cursor = 0
        for sentence in sentences:
            sentence_start = text.find(sentence, cursor)
            if sentence_start < 0:
                sentence_start = cursor
            cursor = sentence_start + len(sentence)
            candidate = f"{buffer}{sentence}" if buffer else sentence
            if count_tokens(candidate) <= chunk_size:
                if not buffer:
                    buffer_start = sentence_start
                buffer = candidate
                continue
            if buffer:
                pieces.append((buffer.strip(), buffer_start))
            buffer = sentence
            buffer_start = sentence_start
        if buffer:
            pieces.append((buffer.strip(), buffer_start))
        if pieces and all(count_tokens(piece) <= chunk_size * 1.15 for piece, _ in pieces):
            return pieces

    return split_by_token_window(text, chunk_size, overlap)


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[。！？.!?；;])", text)
    return [part.strip() for part in parts if part.strip()]


def tail_overlap(text: str, overlap: int) -> str:
    normalized = text.strip()
    if count_tokens(normalized) <= overlap:
        return normalized
    pieces = split_by_token_window(normalized, overlap, 0)
    return pieces[-1][0] if pieces else ""


def count_tokens(text: str) -> int:
    if not text:
        return 0
    encoder = _token_encoder()
    if encoder is not None:
        return len(encoder.encode(text))
    return _fallback_token_count(text)


def split_by_token_window(text: str, chunk_size: int, overlap: int) -> list[tuple[str, int]]:
    units = _token_units(text)
    if not units:
        return []

    pieces: list[tuple[str, int]] = []
    start_index = 0
    step_tokens = max(chunk_size - overlap, 1)
    while start_index < len(units):
        token_total = 0
        end_index = start_index
        while end_index < len(units):
            unit_tokens = max(units[end_index][3], 1)
            if end_index > start_index and token_total + unit_tokens > chunk_size:
                break
            token_total += unit_tokens
            end_index += 1

        start_char = units[start_index][1]
        end_char = units[max(end_index - 1, start_index)][2]
        piece = text[start_char:end_char].strip()
        if piece:
            pieces.append((piece, start_char))

        if end_index >= len(units):
            break
        moved_tokens = 0
        next_start = start_index
        while next_start < len(units) and moved_tokens < step_tokens:
            moved_tokens += max(units[next_start][3], 1)
            next_start += 1
        start_index = max(next_start, start_index + 1)
    return pieces


def _append_chunk(
    *,
    chunks: list[ChunkRecord],
    index: int,
    document_id: str,
    paper_name: str,
    page_start: int,
    page_end: int,
    section: str,
    source: str,
    file_hash: str,
    text: str,
    char_start: int,
    char_end: int,
    chunk_type: str,
) -> int:
    normalized = " ".join(text.split())
    if len(normalized) < 40:
        return index
    chunk_id = f"{document_id}_chunk_{index:05d}"
    chunks.append(
        ChunkRecord(
            chunk_id=chunk_id,
            document_id=document_id,
            paper_name=paper_name,
            page=page_start,
            section=section,
            source=source,
            file_hash=file_hash,
            text=normalized,
            char_start=char_start,
            char_end=char_end,
            quote=normalized[:280],
            page_start=page_start,
            page_end=page_end,
            token_count=count_tokens(normalized),
            chunk_type=chunk_type,
        )
    )
    return index + 1


def _add_neighbor_overlap(chunks: list[ChunkRecord], overlap: int) -> None:
    if overlap <= 0:
        return
    previous_text_chunk: ChunkRecord | None = None
    for chunk in chunks:
        if chunk.chunk_type != "text":
            previous_text_chunk = None
            continue
        if previous_text_chunk and _can_share_context(previous_text_chunk, chunk):
            overlap_text = tail_overlap(previous_text_chunk.text, overlap)
            if overlap_text and not chunk.text.startswith(overlap_text):
                chunk.text = f"{overlap_text}\n\n{chunk.text}"
                chunk.token_count = count_tokens(chunk.text)
        previous_text_chunk = chunk


def _attach_parent_contexts(chunks: list[ChunkRecord], parent_chunk_size: int) -> None:
    if not chunks:
        return
    for index, chunk in enumerate(chunks):
        indices = _parent_window_indices(chunks, index, parent_chunk_size)
        if not indices:
            continue
        parent_chunks = [chunks[item] for item in indices]
        parent_text = "\n\n".join(item.text for item in parent_chunks)
        chunk.parent_id = f"{chunk.document_id}_parent_{indices[0]:05d}_{indices[-1]:05d}"
        chunk.parent_text = parent_text
        chunk.parent_page_start = min(item.page_start for item in parent_chunks)
        chunk.parent_page_end = max(item.page_end for item in parent_chunks)
        chunk.parent_char_start = parent_chunks[0].char_start
        chunk.parent_char_end = parent_chunks[-1].char_end
        chunk.parent_token_count = count_tokens(parent_text)


def _parent_window_indices(chunks: list[ChunkRecord], index: int, token_budget: int) -> list[int]:
    selected = [index]
    token_total = max(chunks[index].token_count, count_tokens(chunks[index].text))
    left = index - 1
    right = index + 1
    while True:
        changed = False
        if right < len(chunks) and _can_share_context(chunks[index], chunks[right]):
            next_tokens = max(chunks[right].token_count, count_tokens(chunks[right].text))
            if token_total + next_tokens <= token_budget:
                selected.append(right)
                token_total += next_tokens
                right += 1
                changed = True
        if left >= 0 and _can_share_context(chunks[left], chunks[index]):
            next_tokens = max(chunks[left].token_count, count_tokens(chunks[left].text))
            if token_total + next_tokens <= token_budget:
                selected.insert(0, left)
                token_total += next_tokens
                left -= 1
                changed = True
        if not changed:
            break
    return selected


def _should_flush_for_block(previous: TextBlock, current: TextBlock) -> bool:
    if previous.section != current.section and current.section != "Unknown":
        return True
    return False


def _can_share_context(left: ChunkRecord, right: ChunkRecord) -> bool:
    if left.document_id != right.document_id:
        return False
    if left.page_end + 1 < right.page_start:
        return False
    if left.section == right.section:
        return True
    return "Unknown" in {left.section, right.section}


def _is_table_context_block(block: TextBlock) -> bool:
    text = " ".join(block.text.split())
    if len(text) > 420:
        return False
    return bool(
        re.search(
            r"\b(table|tab\.)\s*\d+|表\s*[一二三四五六七八九十\d]+|注[:：]|数据来源|资料来源|source",
            text,
            re.IGNORECASE,
        )
    )


@lru_cache(maxsize=1)
def _token_encoder() -> Any | None:
    try:
        import tiktoken
    except Exception:
        return None
    try:
        return tiktoken.get_encoding("cl100k_base")
    except Exception:
        return None


def _fallback_token_count(text: str) -> int:
    total = 0
    for match in TOKEN_UNIT_PATTERN.finditer(text):
        token = match.group(0)
        if token.isspace():
            continue
        if re.fullmatch(r"[\u4e00-\u9fff]", token):
            total += 1
        elif re.fullmatch(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", token):
            total += max(1, math.ceil(len(token) / 4))
        else:
            total += 1
    return total


def _token_units(text: str) -> list[tuple[str, int, int, int]]:
    units: list[tuple[str, int, int, int]] = []
    for match in TOKEN_UNIT_PATTERN.finditer(text):
        value = match.group(0)
        units.append((value, match.start(), match.end(), count_tokens(value)))
    return units
